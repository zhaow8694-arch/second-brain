# -*- coding: utf-8 -*-
"""Generate SniperTrendEA v8.61 overfit verification report (.docx)."""
from pathlib import Path
import csv

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn

OUT = Path(r"E:\grokmacd\SniperTrendEA_v8.61_过拟合验证分析报告.docx")
RESULTS = Path(r"E:\grokmacd\overfit_check\overfit_results.csv")
RANKING = Path(r"E:\grokmacd\overfit_check\robust_ranking.csv")


def set_cell_shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        from docx.oxml import OxmlElement
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "D5E8F0")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def load_csv(path):
    with open(path, encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def main():
    results = load_csv(RESULTS)
    ranking = load_csv(RANKING)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SniperTrendEA v8.61 过拟合验证分析报告")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("品种：XAUUSD · 周期：H4 · 初始资金：20,000 USD").font.size = Pt(11)
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.add_run("报告日期：2026年6月18日").font.size = Pt(10)
    sub2.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_heading("一、分析概述", level=1)
    doc.add_paragraph(
        "本报告对 SniperTrendEA v8.61 共 7 组参数（5 组遗传优化 PASS 参数 + BALANCED + CONSERVATIVE）"
        "在 3 个时段（2015–2019、2020–2025、2025–2026）进行了 21 次样本外/样本内回测验证，"
        "旨在排查遗传优化参数的过拟合风险，并给出实盘参数选择建议。"
    )

    doc.add_heading("二、测试方法与修复说明", level=1)
    bullets = [
        "测试矩阵：7 组参数 × 3 时段 = 21 次回测，品种 XAUUSD H4，每轮顺序启动 MT5 策略测试器。",
        "样本内区间：2020.01.01 – 2025.12.31（与遗传优化目标区间一致）。",
        "样本外区间：2015.01.01 – 2019.12.31（历史验证）、2025.01.01 – 2026.06.30（近期 forward）。",
        "修复项：修正 run_overfit_check.ps1 变量插值 bug（报告名 overfit__ → overfit_{组}_{时段}）；"
        "PASS 参数集改用 BEST_PF 模板格式，确保 MT5 正确加载 InpFilterPreset=3 自定义参数。",
        "第一次无效运行：因 .set 格式错误，6 组 PASS/BALANCED 实际均回退为 EA 默认 BALANCED 参数，已作废并重跑。",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("三、核心发现", level=1)
    doc.add_paragraph(
        "遗传优化报告的盈利因子（PF）无法在相同样本区间内复现，存在严重过拟合。"
        "优化声称 PASS1577 PF=4.05，但相同样本内实测仅 2.15，落差约 47%。"
        "所有优化参数组在 2015–2019 历史样本外表现疲弱（PF 0.96–1.14），"
        "而 2025–2026 近期 PF 偏高（2.57–2.88）但交易数仅 31–33 笔，统计意义不足。"
    )

    doc.add_heading("四、优化声称 PF vs 样本内实测 PF", level=1)
    opt_rows = [
        ("PASS1577", "4.05", "2.15", "-47%"),
        ("PASS1632", "3.90", "2.13", "-45%"),
        ("PASS1729", "3.87", "2.36", "-39%"),
        ("PASS1581", "3.87", "2.06", "-47%"),
        ("PASS1639", "3.72", "1.79", "-52%"),
    ]
    add_table(
        doc,
        ["参数组", "优化声称 PF", "样本内实测 PF (2020-2025)", "落差"],
        opt_rows,
        [3.5, 3.5, 5.5, 2.5],
    )

    doc.add_heading("五、三时段盈利因子矩阵", level=1)
    groups_order = [
        "PASS1577", "PASS1632", "PASS1729", "PASS1581", "PASS1639",
        "BALANCED", "CONSERVATIVE",
    ]
    matrix = {g: {} for g in groups_order}
    for r in results:
        matrix[r["Group"]][r["Period"]] = r["PF"]

    pf_rows = []
    for g in groups_order:
        pf_rows.append((
            g,
            matrix[g].get("2015_2019", "-"),
            matrix[g].get("2020_2025", "-"),
            matrix[g].get("2025_2026", "-"),
        ))
    add_table(
        doc,
        ["参数组", "2015-2019 (OOS)", "2020-2025 (样本内)", "2025-2026 (OOS)"],
        pf_rows,
        [3.5, 3.5, 3.5, 3.5],
    )

    doc.add_heading("六、稳健性排名（样本外加权）", level=1)
    doc.add_paragraph(
        "稳健分 = 样本外 PF × 0.7 + 样本内 PF × 0.3 − 最大回撤% × 0.01。"
        "OOS/样本内比 < 0.85 视为中等过拟合风险。"
    )
    rank_rows = []
    for r in ranking:
        rank_rows.append((
            r["Group"],
            r["PF_2015_2019"],
            r["PF_2020_2025"],
            r["PF_2025_2026"],
            r["OOS_InSample_Ratio"],
            r["MaxDDPct"] + "%",
            r["OverfitFlag"],
            r["RobustScore"],
        ))
    add_table(
        doc,
        ["参数组", "PF 2015-19", "PF 2020-25", "PF 2025-26", "OOS/样本内", "最大回撤", "过拟合标记", "稳健分"],
        rank_rows,
        [2.8, 2.2, 2.2, 2.2, 2.2, 2.2, 2.2, 2.0],
    )

    doc.add_heading("七、过拟合判断", level=1)
    judgments = [
        "2015–2019 样本外：所有优化参数组 PF ≤ 1.14，PASS1729 仅 0.96（亏损），策略对 2020 年前市场环境适配性差。",
        "2020–2025 样本内：优化参数实测 PF 2.0–2.4，远低于优化器报告的 3.7–4.0，确认优化过拟合。",
        "2025–2026 近期：PF 2.57–2.88 看似优异，但仅 31–33 笔交易，样本过小，暂不可作为实盘依据。",
        "最大回撤：优化参数组 52–63%，BALANCED 仅 41%，风险收益比更优。",
        "跨时段稳定性：BALANCED 预设 OOS/样本内比 0.90，三时段 PF 最均衡（1.37 / 1.82 / 1.90）。",
    ]
    for j in judgments:
        doc.add_paragraph(j, style="List Bullet")

    doc.add_heading("八、最终建议", level=1)
    recommendations = [
        "不要直接使用遗传优化 top PF 参数。PASS1577 PF=4.05 是过拟合产物，相同样本内实测仅 2.15，不可作为收益预期。",
        "实盘优先考虑 BALANCED 预设：三时段 PF 最均衡（1.37 / 1.82 / 1.90），最大回撤最低（约 41%），跨时段稳定性最佳。",
        "若坚持使用优化参数：推荐 PASS1577 或 PASS1632，但应按 PF ≈ 2.0 设定预期，而非 4.0；须接受 50%+ 回撤风险。",
        "2025–2026 高 PF 暂不采信：交易笔数过少（31–33 笔），需继续 forward test 至少 3–6 个月后再评估。",
    ]
    for i, rec in enumerate(recommendations, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(rec)

    doc.add_heading("九、附录：数据文件", level=1)
    appendix = [
        r"E:\grokmacd\overfit_check\overfit_results.csv — 21 次回测明细",
        r"E:\grokmacd\overfit_check\robust_ranking.csv — 稳健性排名",
        r"E:\grokmacd\overfit_check\reports\run_summary.csv — 运行摘要",
        r"D:\MT5测试\MetaTrader 5\SingleEAReports\overfit_check_grokmacd_v861\ — HTML 报告目录",
    ]
    for a in appendix:
        doc.add_paragraph(a, style="List Bullet")

    doc.add_paragraph()
    footer = doc.add_paragraph("— 报告由 GROKMACD 过拟合验证流程自动生成 —")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()