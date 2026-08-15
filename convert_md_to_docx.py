import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_styled_table(doc, rows_data, is_header=True):
    """Add a table to the document from markdown rows"""
    if not rows_data:
        return

    max_cols = max(len(row) for row in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=max_cols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows_data):
        for j, cell_text in enumerate(row):
            if j < max_cols:
                cell = table.cell(i, j)
                cell.text = cell_text.strip()

                # Style header row
                if is_header and i == 0:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(9)
                    set_cell_shading(cell, '2F5496')
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                else:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)

    doc.add_paragraph()  # spacing after table

def parse_markdown_table(text):
    """Parse a markdown table into list of rows"""
    lines = text.strip().split('\n')
    rows = []
    for line in lines:
        if '|' in line:
            cells = [c.strip() for c in line.split('|')]
            # Remove empty first/last from split
            if cells and cells[0] == '':
                cells = cells[1:]
            if cells and cells[-1] == '':
                cells = cells[:-1]
            rows.append(cells)

    # Filter out alignment rows (like |:---|:---:|)
    filtered = []
    for row in rows:
        if not all(re.match(r'^:?-{3,}:?$', c) for c in row):
            filtered.append(row)
    return filtered

def convert_markdown_to_docx(md_path, docx_path):
    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Horizontal rule
        if re.match(r'^---+$', line.strip()):
            # Add a horizontal rule-like line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'B0B0B0')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Frontmatter block
        if line.strip() == '---' and i == 0:
            i += 1
            while i < len(lines) and lines[i].strip() != '---':
                i += 1
            i += 1  # skip closing ---
            continue

        # Headers
        header_match = re.match(r'^(#{1,6})\s+(.+?)(?:\s+\{#.*?\})?$', line)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2)
            # Remove markdown links — keep text
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            # Remove bold/italic markers
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'\*([^*]+)\*', r'\1', text)
            text = re.sub(r'`([^`]+)`', r'\1', text)

            if level == 1:
                h = doc.add_heading(text, level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level <= 4:
                doc.add_heading(text, level=level)
            else:
                doc.add_heading(text, level=4)
            i += 1
            continue

        # Code blocks
        if line.strip().startswith('```'):
            lang = line.strip()[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```

            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(50, 50, 50)
                # Add shading
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'F5F5F5')
                shd.set(qn('w:val'), 'clear')
                pPr.append(shd)
            continue

        # Tables (look ahead for multi-line table)
        if '|' in line and i+1 < len(lines) and re.match(r'^\|?\s*:?-{3,}:?\s*\|', lines[i+1]):
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1

            table_text = '\n'.join(table_lines)
            rows = parse_markdown_table(table_text)
            if rows:
                add_styled_table(doc, rows)
            continue

        # Blockquotes
        if line.startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].startswith('>'):
                q_text = lines[i][1:].strip()
                # Clean formatting
                q_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', q_text)
                q_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', q_text)
                quote_lines.append(q_text)
                i += 1

            for q_line in quote_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(q_line)
                run.font.size = Pt(9.5)
                run.font.italic = True
                run.font.color.rgb = RGBColor(80, 80, 80)
                # Left border effect via indent
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                left = OxmlElement('w:left')
                left.set(qn('w:val'), 'single')
                left.set(qn('w:sz'), '12')
                left.set(qn('w:space'), '8')
                left.set(qn('w:color'), '2F5496')
                pBdr.append(left)
                pPr.append(pBdr)
            continue

        # Checklist items
        check_match = re.match(r'^(.*?)-\s+\[([ x])\]\s+(.+)$', line)
        if check_match:
            prefix = check_match.group(1)
            checked = check_match.group(2) == 'x'
            text = check_match.group(3)
            # Clean formatting
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            symbol = '☑' if checked else '☐'
            run = p.add_run(f'{symbol}  {text}')
            run.font.size = Pt(10)
            if checked:
                run.font.color.rgb = RGBColor(0, 128, 0)
            i += 1
            continue

        # Unordered list items
        ul_match = re.match(r'^(\s*)[-*+]\s+(.+)$', line)
        if ul_match:
            indent_level = len(ul_match.group(1)) // 2
            text = ul_match.group(2)
            # Clean formatting / keep bold
            has_bold = '**' in text

            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(1 + indent_level * 0.8)

            if has_bold:
                # Handle bold inline
                parts = re.split(r'(\*\*[^*]+\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        # Clean links
                        part = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', part)
                        part = re.sub(r'`([^`]+)`', r'\1', part)
                        p.add_run(part)
            else:
                text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
                text = re.sub(r'`([^`]+)`', r'\1', text)
                p.add_run(text)

            # Check for sub-items (nested lists on next lines)
            i += 1
            while i < len(lines) and lines[i].strip() and (
                re.match(r'^\s{4,}[-*+]\s+', lines[i]) or
                re.match(r'^\s{4,}\d+[.)]\s+', lines[i]) or
                (lines[i].strip().startswith('└') or lines[i].strip().startswith('├'))
            ):
                sub_text = lines[i].strip()
                sub_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', sub_text)
                sub_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', sub_text)
                sub_text = re.sub(r'`([^`]+)`', r'\1', sub_text)
                sp = doc.add_paragraph(style='List Bullet 2')
                sp.paragraph_format.left_indent = Cm(1.8 + indent_level * 0.8)
                sp.add_run(sub_text)
                i += 1
            continue

        # Ordered list items
        ol_match = re.match(r'^(\s*)\d+[.)]\s+(.+)$', line)
        if ol_match:
            text = ol_match.group(2)
            # Clean bold/links
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)

            p = doc.add_paragraph(style='List Number')
            # Handle bold
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    part = re.sub(r'`([^`]+)`', r'\1', part)
                    p.add_run(part)
            i += 1
            continue

        # "Step" lines with special formatting (like "Step 1: ...")
        step_match = re.match(r'^Step\s+(\d+):\s+(.+)$', line)
        if step_match:
            p = doc.add_paragraph()
            run = p.add_run(f'Step {step_match.group(1)}: ')
            run.bold = True
            run.font.size = Pt(10.5)
            run = p.add_run(step_match.group(2))
            run.font.size = Pt(10.5)
            i += 1
            continue

        # Regular paragraph (non-empty)
        if line.strip():
            text = line.strip()

            # Handle inline formatting
            # Bold
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # We'll style later

            # Links
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)

            # Inline code
            text = re.sub(r'`([^`]+)`', r'\1', text)

            # Emoji — keep as is (Word supports emoji)

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)

            # Re-process bold with actual bold styling
            orig_text = lines[i].strip()
            bold_parts = re.split(r'(\*\*[^*]+\*\*)', orig_text)
            for part in bold_parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.strip():
                    clean = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', part)
                    clean = re.sub(r'`([^`]+)`', r'\1', clean)
                    if clean.strip():
                        p.add_run(clean)

            # Check for continuation lines (not headers, not lists, not empty)
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and not lines[i].startswith('|') and not lines[i].startswith('>') and not lines[i].startswith('- ') and not lines[i].startswith('* ') and not re.match(r'^\d+[.)]', lines[i]) and not lines[i].startswith('```') and not re.match(r'^---+', lines[i].strip()):
                cont_text = lines[i].strip()
                cont_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cont_text)
                cont_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', cont_text)
                cont_text = re.sub(r'`([^`]+)`', r'\1', cont_text)
                if cont_text:
                    run = p.add_run(' ' + cont_text)
                i += 1
            continue

        i += 1

    # Save
    doc.save(docx_path)
    print(f"✅ Word document saved to: {docx_path}")

if __name__ == '__main__':
    import sys
    md_path = sys.argv[1] if len(sys.argv) > 1 else r'E:\知识库\02_市场分析\短剧出海\中亚短剧平台_全流程执行方案_v3.md'
    docx_path = sys.argv[2] if len(sys.argv) > 2 else r'E:\知识库\02_市场分析\短剧出海\中亚短剧平台_全流程执行方案_v3.docx'
    convert_markdown_to_docx(md_path, docx_path)
